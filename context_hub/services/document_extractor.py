"""Text extraction from user-uploaded files.

Accepts ``.md``, ``.txt``, ``.pdf`` and ``.docx`` and returns plain text plus a
suggested title. Heavy parsers (pymupdf, python-docx) are lazy-imported so the
core install does not pull them in.

Failure modes (each raises ``ExtractionError`` with a user-friendly message):
- Unsupported extension
- Required extra not installed
- Parser raises while reading the file
- Resulting text is empty
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import PurePosixPath

_ALLOWED_EXTENSIONS = frozenset({".md", ".txt", ".pdf", ".docx"})


class ExtractionError(Exception):
    """Raised when a file cannot be extracted into usable text."""


@dataclass(frozen=True)
class ExtractedDocument:
    title: str
    text: str
    extension: str  # lowercase, includes the leading dot


def supported_extensions() -> frozenset[str]:
    return _ALLOWED_EXTENSIONS


def extract(filename: str, data: bytes) -> ExtractedDocument:
    """Extract plain text from a file payload.

    Args:
        filename: Original filename (used for extension detection and title).
        data:     Raw file bytes.

    Raises:
        ExtractionError: when extraction fails or the file is unsupported.
    """
    if not filename:
        raise ExtractionError("Filename is required to determine file type.")
    if not data:
        raise ExtractionError("Uploaded file is empty.")

    stem = PurePosixPath(filename).stem
    extension = PurePosixPath(filename).suffix.lower()
    if extension not in _ALLOWED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file extension '{extension}'. "
            f"Allowed: {', '.join(sorted(_ALLOWED_EXTENSIONS))}."
        )

    if extension in (".md", ".txt"):
        text = _decode_text(data)
    elif extension == ".pdf":
        text = _extract_pdf(data)
    else:
        text = _extract_docx(data)

    text = text.strip()
    if not text:
        raise ExtractionError(
            f"No extractable text found in '{filename}'. "
            "Scanned PDFs without OCR cannot be ingested."
        )

    title = _detect_title(text, fallback=stem) if extension == ".md" else stem
    return ExtractedDocument(title=title, text=text, extension=extension)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp932", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError(
        "Could not decode the file as text. Save it as UTF-8 and retry."
    )


def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ExtractionError(
            "PDF support requires the [documents] extra. "
            "Install with: pip install 'yohakuforce-context-hub[documents]'"
        ) from exc
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            pages = [page.get_text("text") for page in doc]
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Failed to parse PDF: {exc}") from exc
    return "\n\n".join(p.strip() for p in pages if p and p.strip())


def _extract_docx(data: bytes) -> str:
    try:
        import io

        import docx  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ExtractionError(
            "DOCX support requires the [documents] extra. "
            "Install with: pip install 'yohakuforce-context-hub[documents]'"
        ) from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Failed to parse DOCX: {exc}") from exc

    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                chunks.append(row_text)
    return "\n\n".join(chunks)


def _detect_title(markdown_text: str, fallback: str) -> str:
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            return stripped[2:].strip() or fallback
        break
    return fallback
