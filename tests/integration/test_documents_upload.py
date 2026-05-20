"""Integration tests for POST /api/v1/documents/upload.

Exercises both the unit-level extractor (markdown / txt / pdf / docx) and the
end-to-end multipart route. PDF and DOCX tests require the [documents] extra to
be installed; otherwise they skip.
"""

from __future__ import annotations

import io

import pytest
from httpx import ASGITransport, AsyncClient

from context_hub.services.document_extractor import (
    ExtractionError,
    extract,
    supported_extensions,
)
from context_hub.shared.types import ProjectId, SourceType

from tests.integration.test_api_routers import (  # type: ignore[attr-defined]
    _HEADERS,
    app,
    repos,
)

# Re-export fixtures so pytest finds them in this module's namespace.
__all__ = ["app", "repos"]


# ---------------------------------------------------------------------------
# Extractor unit coverage
# ---------------------------------------------------------------------------


class TestDocumentExtractor:
    def test_supported_extensions_contains_expected_set(self):
        assert supported_extensions() == frozenset({".md", ".txt", ".pdf", ".docx"})

    def test_md_extraction_uses_first_h1_as_title(self):
        data = b"# Kickoff\n\nbody text here"
        out = extract("kickoff.md", data)
        assert out.title == "Kickoff"
        assert "body text here" in out.text
        assert out.extension == ".md"

    def test_md_falls_back_to_filename_when_no_h1(self):
        out = extract("plain.md", b"no heading here")
        assert out.title == "plain"

    def test_txt_uses_filename_stem_as_title(self):
        out = extract("notes.txt", b"some text")
        assert out.title == "notes"
        assert out.text == "some text"

    def test_unsupported_extension_raises(self):
        with pytest.raises(ExtractionError, match="Unsupported file extension"):
            extract("slides.pptx", b"binary")

    def test_empty_file_raises(self):
        with pytest.raises(ExtractionError, match="empty"):
            extract("x.md", b"")

    def test_missing_filename_raises(self):
        with pytest.raises(ExtractionError, match="Filename is required"):
            extract("", b"body")

    def test_whitespace_only_raises_no_extractable_text(self):
        with pytest.raises(ExtractionError, match="No extractable text"):
            extract("blank.md", b"   \n   \n  ")

    def test_pdf_extraction(self):
        pytest.importorskip("fitz")
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "PDF body content for context-hub")
        data = doc.tobytes()
        doc.close()

        out = extract("report.pdf", data)
        assert "PDF body content" in out.text
        assert out.title == "report"

    def test_docx_extraction(self):
        pytest.importorskip("docx")
        import docx

        d = docx.Document()
        d.add_paragraph("First paragraph.")
        d.add_paragraph("Second paragraph with details.")
        buf = io.BytesIO()
        d.save(buf)

        out = extract("memo.docx", buf.getvalue())
        assert "First paragraph." in out.text
        assert "Second paragraph" in out.text
        assert out.title == "memo"


# ---------------------------------------------------------------------------
# End-to-end multipart route
# ---------------------------------------------------------------------------


class TestDocumentUploadEndpoint:
    @pytest.mark.asyncio
    async def test_md_upload_returns_201_and_persists(self, app, repos):
        _job, doc_repo, _issue, _proj, _emb = repos
        files = {"file": ("plan.md", b"# Plan\n\nWork breakdown.", "text/markdown")}
        data = {
            "project_id": "proj-001",
            "source_type": "meeting",
        }
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 201, resp.text
        body = resp.json()
        assert body["success"] is True
        assert body["data"]["external_id"] == "plan.md"

        stored = await doc_repo.find_by_external_id(
            ProjectId("proj-001"), SourceType.MEETING, "plan.md"
        )
        assert stored is not None
        assert "Work breakdown." in stored.raw_content.text
        assert stored.raw_content.text.startswith("# Plan")

    @pytest.mark.asyncio
    async def test_txt_upload_uses_filename_as_title(self, app, repos):
        _job, doc_repo, *_ = repos
        files = {"file": ("memo.txt", b"plain memo body", "text/plain")}
        data = {"project_id": "proj-001", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 201
        stored = await doc_repo.find_by_external_id(
            ProjectId("proj-001"), SourceType.FILE, "memo.txt"
        )
        assert stored is not None
        assert stored.raw_content.text.startswith("# memo\n")

    @pytest.mark.asyncio
    async def test_upload_explicit_external_id_overrides_filename(self, app, repos):
        _job, doc_repo, *_ = repos
        files = {"file": ("doc.md", b"# T\n\nbody", "text/markdown")}
        data = {
            "project_id": "proj-001",
            "source_type": "file",
            "external_id": "stable-id-v1",
        }
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 201
        assert resp.json()["data"]["external_id"] == "stable-id-v1"

    @pytest.mark.asyncio
    async def test_upload_rejects_unsupported_extension(self, app):
        files = {"file": ("slides.pptx", b"binary blob", "application/octet-stream")}
        data = {"project_id": "proj-001", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 400
        assert "Unsupported file extension" in resp.text

    @pytest.mark.asyncio
    async def test_upload_rejects_unknown_project(self, app):
        files = {"file": ("plan.md", b"# Plan\n\nbody", "text/markdown")}
        data = {"project_id": "no-such-project", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 404

    @pytest.mark.asyncio
    async def test_upload_requires_write_scope(self, app):
        files = {"file": ("plan.md", b"# Plan\n\nbody", "text/markdown")}
        data = {"project_id": "proj-001", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files
            )
        assert resp.status_code == 401

    @pytest.mark.asyncio
    async def test_upload_rejects_oversized_file(self, app):
        big = b"a" * (10 * 1024 * 1024 + 1)
        files = {"file": ("big.txt", big, "text/plain")}
        data = {"project_id": "proj-001", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 413

    @pytest.mark.asyncio
    async def test_upload_pdf_round_trip(self, app, repos):
        pytest.importorskip("fitz")
        import fitz

        _job, doc_repo, *_ = repos
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Contract terms version 1.2")
        pdf_bytes = doc.tobytes()
        doc.close()

        files = {"file": ("contract.pdf", pdf_bytes, "application/pdf")}
        data = {"project_id": "proj-001", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 201
        stored = await doc_repo.find_by_external_id(
            ProjectId("proj-001"), SourceType.FILE, "contract.pdf"
        )
        assert stored is not None
        assert "Contract terms version 1.2" in stored.raw_content.text

    @pytest.mark.asyncio
    async def test_upload_docx_round_trip(self, app, repos):
        pytest.importorskip("docx")
        import docx

        _job, doc_repo, *_ = repos
        d = docx.Document()
        d.add_paragraph("見積書の内訳")
        d.add_paragraph("人月単価 100 万円")
        buf = io.BytesIO()
        d.save(buf)

        files = {
            "file": (
                "estimate.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        data = {"project_id": "proj-001", "source_type": "file"}
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/documents/upload", data=data, files=files, headers=_HEADERS
            )
        assert resp.status_code == 201
        stored = await doc_repo.find_by_external_id(
            ProjectId("proj-001"), SourceType.FILE, "estimate.docx"
        )
        assert stored is not None
        assert "見積書の内訳" in stored.raw_content.text
        assert "人月単価" in stored.raw_content.text

    @pytest.mark.asyncio
    async def test_supported_extensions_endpoint(self, app):
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.get(
                "/api/v1/documents/upload/supported-extensions",
                headers=_HEADERS,
            )
        assert resp.status_code == 200
        body = resp.json()
        assert ".md" in body["data"]["extensions"]
        assert ".txt" in body["data"]["extensions"]
        assert ".pdf" in body["data"]["extensions"]
        assert ".docx" in body["data"]["extensions"]
